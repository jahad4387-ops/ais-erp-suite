import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """设置单元格背景颜色"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """设置单元格边距（单位：dxa，1 pt = 20 dxa）"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    设置单元格边框
    kwargs: top, bottom, left, right, insideH, insideV
    value should be dict: {'sz': 12, 'val': 'single', 'color': 'D3D3D3'}
    """
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        if border_name in kwargs:
            edge = OxmlElement(f'w:{border_name}')
            for key, val in kwargs[border_name].items():
                edge.set(qn(f'w:{key}'), str(val))
            tcBorders.append(edge)
    tcPr.append(tcBorders)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    """添加带样式的标题，确保字体为微软雅黑"""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    # 设置颜色与字体
    run = p.runs[0]
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(24, 144, 255) # 经典AntD蓝色
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 31, 31)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(77, 77, 77)
    return p

def add_paragraph_styled(doc, text="", bold_prefix=None, space_after=6, line_spacing=1.15):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Microsoft YaHei'
        r_bold._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r_bold.font.size = Pt(10.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(51, 51, 51)
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = 'Microsoft YaHei'
        r_text._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r_text.font.size = Pt(10.5)
        r_text.font.color.rgb = RGBColor(77, 77, 77)
        
    return p

def add_bullet_styled(doc, text, bold_prefix=None):
    """添加列表项"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Microsoft YaHei'
        r_bold._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r_bold.font.size = Pt(10.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(51, 51, 51)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Microsoft YaHei'
    r_text._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = RGBColor(77, 77, 77)
    return p

def create_table_styled(doc, headers, col_widths=None):
    """创建美观的表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 格式化表头
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], '1890FF') # 经典AntD蓝色背景
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
        # 居中和字体设置
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # 设置边框
        border_spec = {'sz': 4, 'val': 'single', 'color': 'D3D3D3'}
        set_cell_border(hdr_cells[i], 
                        top=border_spec, bottom=border_spec, 
                        left=border_spec, right=border_spec)
        
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = Inches(col_widths[i])
            
    return table

def add_table_row_styled(table, row_data, col_widths=None, is_even=False):
    """向表格添加一行，并设置样式"""
    row = table.add_row()
    cells = row.cells
    bg_color = 'F5F5F5' if is_even else 'FFFFFF' # 斑马线
    
    for i, text in enumerate(row_data):
        cells[i].text = str(text)
        set_cell_background(cells[i], bg_color)
        set_cell_margins(cells[i], top=100, bottom=100, left=150, right=150)
        
        p = cells[i].paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        if len(p.runs) > 0:
            run = p.runs[0]
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(51, 51, 51)
            
        # 设置边框
        border_spec = {'sz': 4, 'val': 'single', 'color': 'E8E8E8'}
        set_cell_border(cells[i], 
                        top=border_spec, bottom=border_spec, 
                        left=border_spec, right=border_spec)
        
        if col_widths and i < len(col_widths):
            cells[i].width = Inches(col_widths[i])
            
    return row

def generate_manual():
    doc = docx.Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- 1. 封面 ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(10)
    
    run_title = title_p.add_run("AIS-ERP-Suite 智能财务管理系统")
    run_title.font.name = 'Microsoft YaHei'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(24, 144, 255)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(200)
    run_sub = subtitle_p.add_run("用户使用说明手册 (Phase 1 - Phase 5)")
    run_sub.font.name = 'Microsoft YaHei'
    run_sub._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(102, 102, 102)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = info_p.add_run("系统版本：Enterprise v1.5.0\n文档状态：已完成 Phase 1-5 模块交付\n发布日期：2026年6月")
    run_info.font.name = 'Microsoft YaHei'
    run_info._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = RGBColor(128, 128, 128)
    
    # 插入分页符进入目录与正文
    doc.add_page_break()
    
    # --- 2. 目录说明/前言 ---
    add_heading_styled(doc, "前言与系统概述", level=1)
    add_paragraph_styled(doc, "AIS-ERP-Suite 是一款面向中国中小企业（SMEs）设计的高端、AI 原生财务与业务一体化管理系统。系统严格遵守《企业会计准则（CAS）》，提供跨账套隔离的多主体管理能力，并采用先进的“智能助手（Agent-Native）”架构，允许 AI 助手在人类授权的前提下协同完成财务记账与分析。")
    
    add_heading_styled(doc, "产品核心设计原则", level=2)
    add_bullet_styled(doc, "：系统所有的业务与财务操作均通过标准的 OpenAPI 暴露给 AI 助手，使其能高效辅助数据录入与报表生成工作。", "AI-Native (智能原生)")
    add_bullet_styled(doc, "：所有资金、库存、资产与凭证的终态操作均要求人类用户进行显式确认与审核，防范AI自主操作带来的合规风险。", "Human-in-the-Loop (人机协同)")
    add_bullet_styled(doc, "：数据以“账套”为核心单位进行完全隔离，确保不同子公司、不同独立核算实体的安全性与合规性。", "账套绝对隔离")
    add_bullet_styled(doc, "：系统对一切状态变更操作引入 Idempotency-Key 校验，并对修改和删除数据保留完整的操作轨迹审计日志，不允许物理删除任何财务历史明细。", "数据不可篡改与全面审计")
    
    # --- 3. Phase 1: 总账与系统基础配置 ---
    add_heading_styled(doc, "第一章：总账与系统基础配置 (Phase 1)", level=1)
    add_paragraph_styled(doc, "本章介绍系统平台初始化、基础科目与账套的创建、以及会计期间与初始余额的管理。这是整个系统的财务基石。")
    
    add_heading_styled(doc, "1.1 登录与认证", level=2)
    add_paragraph_styled(doc, "系统采用基于 JWT (JSON Web Token) 的安全访问控制。默认系统管理员账号为 admin，初始密码为 admin123（首次登录后建议立即修改）。用户登录后将获取与账号和账套绑定的临时令牌，并在侧边栏或顶部导航选择目标工作账套。")
    
    add_heading_styled(doc, "1.2 账套管理 (Account Set)", level=2)
    add_paragraph_styled(doc, "账套是独立的业务和财务核算边界，各账套之间的数据完全隔离。")
    add_bullet_styled(doc, "点击“平台管理” -> “账套管理”，进入账套列表。", "创建账套")
    add_bullet_styled(doc, "点击“新建账套”，填写账套编码、账套名称、本位币（默认人民币 CNY）、会计年度开始月份（中国通常为1月）。", "参数配置")
    add_bullet_styled(doc, "新建账套后，系统会自动创建基础表结构，并隔离在专用的数据容器中。", "数据隔离")
    
    add_heading_styled(doc, "1.3 用户、角色与权限管理 (RBAC)", level=2)
    add_paragraph_styled(doc, "系统提供基于角色的权限控制（RBAC），支持多账套交叉授权：")
    add_bullet_styled(doc, "针对特定用户，可授予其在特定账套下的操作权限（如“账套 A 的制单员，账套 B 的只读审核员”）。", "跨账套授权")
    add_bullet_styled(doc, "财务核心权限分离，包括：制单人（Maker）、审核人（Reviewer）、过账人（Poster）。同一个账号在单笔凭证中不能身兼多职。", "权限隔离 (Segregation of Duties)")
    
    add_heading_styled(doc, "1.4 会计科目与编码规则", level=2)
    add_paragraph_styled(doc, "系统支持多达4级会计科目树结构：")
    add_bullet_styled(doc, "科目类型分为：资产、负债、所有者权益、成本、共同、损益。科目余额方向包括借方（Debit）和贷方（Credit）。", "科目分类")
    add_bullet_styled(doc, "每个科目可以配置为是否需要辅助核算（如往来单位、部门、项目等）。", "辅助核算标志")
    add_bullet_styled(doc, "支持导入标准会计科目表（CSV模板）。模板包含字段：科目编码、科目名称、科目类型、余额方向、是否允许手工录入、是否需要辅助核算。", "科目导入")
    
    # 科目类别表格
    t_acc = create_table_styled(doc, ["科目类别", "余额方向示例", "对应业务范畴"], [2.0, 2.0, 3.5])
    add_table_row_styled(t_acc, ["资产 (Asset)", "借方 (Debit)", "库存现金、银行存款、应收账款、固定资产"], is_even=False)
    add_table_row_styled(t_acc, ["负债 (Liability)", "贷方 (Credit)", "应付账款、短期借款、应交税费"], is_even=True)
    add_table_row_styled(t_acc, ["所有者权益 (Equity)", "贷方 (Credit)", "实收资本、盈余公积、本年利润"], is_even=False)
    add_table_row_styled(t_acc, ["成本 (Cost)", "借方 (Debit)", "生产成本、制造费用、劳务成本"], is_even=True)
    add_table_row_styled(t_acc, ["损益 (Revenue/Expense)", "借方/贷方", "主营业务收入（贷）、管理费用（借）"], is_even=False)
    
    doc.add_paragraph() # 空行隔开
    
    add_heading_styled(doc, "1.5 会计期间管理", level=2)
    add_paragraph_styled(doc, "会计期间状态流转为：未启用 -> 开启 (Open) -> 关闭 (Closed) -> 锁定 (Locked)。")
    add_bullet_styled(doc, "只有处于“开启”状态的期间才允许新建和录入会计凭证。", "凭证控制")
    add_bullet_styled(doc, "期末结账后，期间应变更为“关闭”状态，以防止意外修改。如果发现错误，在未锁定的情况下可以重新开启。", "期末关闭")
    add_bullet_styled(doc, "“锁定”是不可逆的永久操作。锁定后任何人都无法重新开启该期间或修改其中的任何财务凭证。", "永久锁定")
    
    add_heading_styled(doc, "1.6 辅助核算与期初余额", level=2)
    add_bullet_styled(doc, "系统提供5个预置辅助核算维度：部门 (Department)、员工 (Employee)、客户 (Customer)、供应商 (Supplier)、项目 (Project)。辅助核算项在凭证分录行中可进行细粒度挂接。", "辅助核算")
    add_bullet_styled(doc, "在首个会计期间启用前，用户必须输入各科目的期初余额，对于要求辅助核算的科目，期初余额须按辅助维度细分。录入完毕后需通过“试算平衡”校验（资产=负债+所有者权益），平衡后方可开始日常核算。", "期初余额")
    
    add_heading_styled(doc, "1.7 凭证管理与核心审批流", level=2)
    add_paragraph_styled(doc, "凭证分为“记账凭证 (J)”、“收款凭证 (R)”、“付款凭证 (P)”和“转账凭证 (T)”。其核心审批和流转步骤如下：")
    add_bullet_styled(doc, "：填制凭证，保存为草稿。系统会自动实时验证：1) 借贷金额必须相等；2) 期间必须处于“开启”状态；3) 科目必须允许手工录入，且已填报必填的辅助核算维度；4) 单笔凭证借贷方不为零。", "1. 制单 (Draft)")
    add_bullet_styled(doc, "：制单员保存凭证后，点击提交送审。凭证状态变为“已提交 (Submitted)”。", "2. 提交 (Submit)")
    add_bullet_styled(doc, "：具有审核权限的独立人员对凭证的借贷科目、附件进行复核，审核通过后状态变更为“已审核 (Approved)”。若发现问题可驳回（Rejected）至草稿状态重新编辑。", "3. 审核 (Approve)")
    add_bullet_styled(doc, "：过账员对已审核的凭证进行批处理过账。过账后状态变更为“已过账 (Posted)”，并正式计入总分类账及辅助明细账，自动更新期末余额。", "4. 过账 (Post)")
    
    add_heading_styled(doc, "1.8 银行对账", level=2)
    add_paragraph_styled(doc, "系统支持银行对账单导入与流水勾对：")
    add_bullet_styled(doc, "支持导入标准银行对账单 CSV 文件。", "对账单导入")
    add_bullet_styled(doc, "支持配置自动对账规则，可按照“金额相同，交易日期在±3天内”等条件自动进行勾对。", "自动核对")
    add_bullet_styled(doc, "手动对账界面允许用户对差异流水进行手工连线核销，并输出标准的银行存款余额调节表，显示银行对账单余额、银行未付企业已付、企业未收银行已收等调节项。", "余额调节表")
    
    add_heading_styled(doc, "1.9 智能凭证助手 (AI Assistant)", level=2)
    add_paragraph_styled(doc, "AI 助手是系统极具特色的亮点模块，允许用户使用自然语言或上传发票/合同文件快速生成凭证：")
    add_bullet_styled(doc, "用户可以输入“用工行网银支付了行政部门租金 5000 元”，或者上传发票 PDF/图片。", "自然语言或文件输入")
    add_bullet_styled(doc, "AI 解析交易意图，自动匹配推荐会计科目（如：借 管理费用-房租 5000 / 贷 银行存款-工行 5000），并生成凭证草稿推荐。", "智能推荐 (Dry Run)")
    add_bullet_styled(doc, "AI 生成的内容永远以“草稿推荐”形式展现。必须由人类会计师审核、确认并点击“转换为正式凭证”后，才会进入系统凭证工作流。AI 无权直接提交或审核凭证。", "确认机制")
    
    # --- 4. Phase 2: 采购与销售 / 应收应付管理 ---
    add_heading_styled(doc, "第二章：采购与销售 / 应收应付管理 (Phase 2)", level=1)
    add_paragraph_styled(doc, "本章介绍客商（供应商和客户）主数据管理、采购与销售业务链条、以及往来核销与账龄分析。")
    
    add_heading_styled(doc, "2.1 合作伙伴主数据 (Partner)", level=2)
    add_paragraph_styled(doc, "系统将供应商与客户统一抽象为“合作伙伴 (Partner)”：")
    add_bullet_styled(doc, "往来单位可以是“客户”、“供应商”，或者同时具备“双重身份”。", "身份属性")
    add_bullet_styled(doc, "包含名称、统一社会信用代码（税号）、银行账号、地址及联系人信息。", "基本档案")
    add_bullet_styled(doc, "配置结算条款（如月结30天）、默认税率（如13%）、默认结算方式（网银/电汇）和信用额度（Credit Limit）。", "核心控制参数")
    
    add_heading_styled(doc, "2.2 采购业务与应付账款 (Procure-to-Pay)", level=2)
    add_paragraph_styled(doc, "系统实现采购全链路闭环核算：")
    add_bullet_styled(doc, "录入采购合同与订单，指定供应商、商品明细与约定单价，跟踪订单审批状态。", "1. 采购订单 (PO)")
    add_bullet_styled(doc, "库管员核对实物数量，做入库登记，生成采购入库单。", "2. 采购入库 (Receipt)")
    add_bullet_styled(doc, "收到供应商发票后，录入系统并与入库单进行勾稽核对，生成应付账款（AP）。系统会自动生成对应的记账凭证草稿（借：原材料 / 贷：应付账款）。", "3. 采购发票与应付确认 (Invoice)")
    add_bullet_styled(doc, "根据付款条件生成付款申请单，审批通过后，记录出纳付款，冲销应付账款，并生成付款凭证草稿（借：应付账款 / 贷：银行存款）。", "4. 付款申请与供应商付款 (Payment)")
    
    add_heading_styled(doc, "2.3 销售业务与应收账款 (Order-to-Cash)", level=2)
    add_paragraph_styled(doc, "销售流程与采购流程镜像对称：")
    add_bullet_styled(doc, "登记与客户签署的销售合同或销售订单，校验客户信用额度。", "1. 销售订单 (SO)")
    add_bullet_styled(doc, "发货后做仓库出库确认，生成销售出库单。", "2. 销售出库 (Delivery)")
    add_bullet_styled(doc, "向客户开具发票时，确认应收账款（AR），并自动生成销售记账凭证草稿（借：应收账款 / 贷：主营业务收入、应交税费-应交增值税-销项税额）。", "3. 销售开票与应收确认 (Invoice)")
    add_bullet_styled(doc, "银行收到货款后录入收款单，冲销客户应收账款余额，自动生成收款凭证草稿。", "4. 客户收款 (Receipt)")
    
    add_heading_styled(doc, "2.4 账龄分析与核销管理", level=2)
    add_paragraph_styled(doc, "系统提供强大的往来报表分析和销账功能：")
    add_bullet_styled(doc, "支持对多发票与多收付款单进行交叉局部核销（如一张大额收款单拆分核销三个销售发票）。", "多对多核销")
    add_bullet_styled(doc, "提供 0-30天、31-60天、61-90天、90天以上等时间区间段的应收/应付账龄分析表（Aging Report），并对超期账款进行信用预警，以便进行资金回笼催收。", "账龄分析")
    
    # --- 5. Phase 3: 存货、生产与成本分摊 ---
    add_heading_styled(doc, "第三章：存货、生产与成本分摊 (Phase 3)", level=1)
    add_paragraph_styled(doc, "本章介绍存货主数据、BOM管理、多仓储管理、生产工单与成本核算引擎。")
    
    add_heading_styled(doc, "3.1 存货基础档案与期初", level=2)
    add_bullet_styled(doc, "存货类别分为：原材料、半成品、产成品、包装物。计量单位支持主副单位自动换算。", "存货属性")
    add_bullet_styled(doc, "每个存货项目均需绑定成本计价方法：先进先出法 (FIFO)、加权平均法 (Moving Average)、标准成本法 (Standard Cost)。系统支持配置批次追踪（Batch）和序列号追踪（Serial）属性。", "成本核算属性")
    add_bullet_styled(doc, "录入存货期初数量、期初单价和总额。可按仓库、货位、批次维度拆分输入。存货期初试算平衡后，才允许进行后续出入库业务。", "存货期初")
    
    add_heading_styled(doc, "3.2 BOM (物料清单) 管理", level=2)
    add_paragraph_styled(doc, "支持建立多级物料清单（BOM），定义生产一单位产成品所需的各原料、半成品的标准配比。支持配置工艺废品率（Scrap Rate）、生产产出率（Yield Rate）和BOM版本管理，为生产领料及标准成本分摊打下基础。")
    
    add_heading_styled(doc, "3.3 仓库、货位与库内流转", level=2)
    add_paragraph_styled(doc, "系统支持多仓库、多货位（Bin）精细管理：")
    add_bullet_styled(doc, "除常规的采购入库、销售出库外，系统支持领料出库、产成品入库、盘盈盘亏等出入库事务类型。", "常规入出库")
    add_bullet_styled(doc, "记录存货在不同仓库或货位之间的转移流转，保持数量和成本的同步流转。", "库间调拨 (Transfer)")
    add_bullet_styled(doc, "支持周期性盘点。系统自动生成盘点报告草稿，比对系统数量与实盘数量并计算差异，经审核后自动产生盘盈盘亏库存调整，并生成对应的财务调整凭证草稿。", "实物盘点 (Stock Count)")
    
    add_heading_styled(doc, "3.4 生产工单与领料完工", level=2)
    add_paragraph_styled(doc, "生产制造核心流转如下：")
    add_bullet_styled(doc, "依据销售预测或销售订单创建生产工单（Work Order），工单状态包含：草稿 -> 释放 -> 生产中 -> 完工 -> 关闭。", "1. 生产工单 (WO)")
    add_bullet_styled(doc, "根据工单对应的BOM配比，生成物料领料单（Material Requisition），从仓库发出原材料进入“生产成本-直接材料”科目。", "2. 生产领料")
    add_bullet_styled(doc, "产品生产完工后，做产品收货（Product Receipt），产成品增加，生产工单标记为完工状态，归集直接材料成本。", "3. 完工入库")
    
    add_heading_styled(doc, "3.5 成本核算与分摊引擎", level=2)
    add_paragraph_styled(doc, "本模块是 Phase 3 的核心难点，确保制造企业的成本真实归集：")
    add_bullet_styled(doc, "：用户可通过 Mock Cost Input 录入生产线相关的直接人工（Direct Labor）以及制造费用（Overhead，如电费、折旧费）。（在 Phase 4/5 启用后，这些折旧与薪资数据将从对应模块直接拉取，不再依赖手动录入）。", "1. 成本费用录入")
    add_bullet_styled(doc, "：成本分摊引擎允许用户进行“分摊试算 (Dry-run)”。根据设定的分配基础（如工时、产量配比），将当期直接人工和制造费用分配给各个生产工单。", "2. 分摊引擎试算与锁定")
    add_bullet_styled(doc, "：试算无误后执行“正式分摊”，系统计算出各完工产成品的最终单位成本，更新存货明细账的成本价格。对零数量存货进行残值调整结转（Zero-quantity residual adjustment），确保单价与金额计算的绝对精确。", "3. 成本结转与残值结清")
    add_bullet_styled(doc, "：生成成本分摊的正式财务记账凭证（借：生产成本-直接人工/制造费用，贷：应付职工薪酬/累计折旧；以及结转至 产成品 的凭证草稿），并进行库存明细账与总账余额的自动对账，提示差异。", "4. 凭证生成与库存对账")
    
    # --- 6. Phase 4: 固定资产与折旧管理 ---
    add_heading_styled(doc, "第四章：固定资产与折旧管理 (Phase 4)", level=1)
    add_paragraph_styled(doc, "本章介绍固定资产折旧方法、资产卡片维护、折旧计算与固定资产处置流程。")
    
    add_heading_styled(doc, "4.1 资产类别与折旧设置", level=2)
    add_paragraph_styled(doc, "在增加固定资产前，需要先创建“资产类别 (Asset Category)”：")
    add_bullet_styled(doc, "如机器设备、电子设备、房屋建筑物、运输工具等。", "类别定义")
    add_bullet_styled(doc, "设定类别的默认预计使用年限、预计净残值率、以及折旧方法。", "默认折旧参数")
    add_bullet_styled(doc, "配置对应资产科目、累计折旧科目、折旧费用科目、处置损益科目等，以便折旧时自动生成正确的会计分录。", "科目映射表")
    
    add_heading_styled(doc, "4.2 固定资产卡片管理 (Asset Card)", level=2)
    add_paragraph_styled(doc, "每一项固定资产在系统中拥有一张独立的固定资产卡片，包含字段：资产编号、资产名称、规格型号、购入日期、启用日期、原值、预计残值、折旧年限、已提折旧、折旧状态、使用部门、存放地点等。")
    add_paragraph_styled(doc, "资产状态包含：草稿 -> 正常折旧 -> 提足折旧 -> 处置销账。卡片支持变更原值、折旧年限或使用部门（记录折旧费用分配部门的变更历史）。")
    
    add_heading_styled(doc, "4.3 折旧计提算法与月度折旧计提", level=2)
    add_paragraph_styled(doc, "系统预置四种主流折旧方法，完全符合 CAS 规范：")
    add_bullet_styled(doc, "：年折旧额 = (原值 - 预计净残值) / 预计使用年限。每月平均计提。", "1. 平均年限法 (Straight-line)")
    add_bullet_styled(doc, "：年折旧率 = 2 / 预计使用年限；折旧额 = 账面净值 * 折旧率。最后两年改为直线法。", "2. 双倍余额递减法 (Double Declining)")
    add_bullet_styled(doc, "：折旧率 = (预计使用年限 - 已使用年数) / [预计使用年限 * (预计使用年限 + 1) / 2]；折旧额 = (原值 - 预计净残值) * 折旧率。", "3. 年数总和法 (Sum-of-the-years-digits)")
    add_bullet_styled(doc, "：根据月度实际工作量（如行驶里程、生产件数）占总预计工作量的配比计算当期折旧。", "4. 工作量法 (Units of Production)")
    
    add_paragraph_styled(doc, "：每月期末，折旧计提引擎批量计算当月所有资产的折旧额。计算出的折旧数额会根据资产卡片指定的“使用部门”自动分摊到对应的“管理费用”、“销售费用”或“制造费用”科目，生成折旧记账凭证草稿（借：管理费用/制造费用-折旧，贷：累计折旧）。", "月度折旧计提流程")
    add_paragraph_styled(doc, "：计提生成的折旧数额中归属于“制造费用”的部分，会自动作为“折旧费用”真实输入项直接同步到 Phase 3 的成本分摊引擎中，实现业务系统的成本无缝贯通。", "Phase 3 成本引擎集成")
    
    add_heading_styled(doc, "4.4 资产处置与盘点", level=2)
    add_bullet_styled(doc, "当资产报废、出售或毁损时，在系统录入资产处置单。系统自动计算处置日期的账面原值、已提折旧和账面净值，根据收回的残料变价或出售收入，计算处置利得或损失，并生成处置记账凭证（借：固定资产清理、累计折旧，贷：固定资产；并结转清理损益至 营业外收支）。", "资产处置 (Disposal)")
    add_bullet_styled(doc, "支持固定资产条码/二维码盘点。生成资产盘点表，录入实盘结果，记录盘盈盘亏资产，并进行相应的账务调整。", "资产盘点 (Asset Count)")
    add_bullet_styled(doc, "提供固定资产增减变动表、折旧明细表、固定资产卡片登记簿以及固定资产与总账对账表，核对资产卡片原值/累计折旧合计数与总账相应科目余额是否一致。", "对账与报表")
    
    # --- 7. Phase 5: 报表管理与经营多维分析 ---
    add_heading_styled(doc, "第五章：报表管理与经营多维分析 (Phase 5)", level=1)
    add_paragraph_styled(doc, "本章介绍报表模板配置、UFO报表设计器的公式编辑、预算管理以及高级多维分析功能。")
    
    add_heading_styled(doc, "5.1 报表模板管理与标准财务报表", level=2)
    add_paragraph_styled(doc, "系统预置符合财政部标准格式的四大财务报表：")
    add_bullet_styled(doc, "：反映企业在特定日期（期末）的资产、负债和所有者权益状况。数据直接根据总账科目期末余额计算填列。", "1. 资产负债表 (Balance Sheet)")
    add_bullet_styled(doc, "：反映企业在一定会计期间的经营成果（收入、费用、利润）。数据根据当期损益类科目的发生额统计填列。", "2. 利润表 (Income Statement)")
    add_bullet_styled(doc, "：反映企业一定期间内现金和现金等价物的流入和流出。支持直接法，通过标记银行存款收支凭证的现金流量项目进行动态统计。", "3. 现金流量表 (Cash Flow Statement)")
    add_bullet_styled(doc, "：用户可在“报表模板”页面自定义报表行次、名称，并配置计算公式。报表支持多版本升级与历史快照归档。", "4. 自定义报表模板")
    
    add_heading_styled(doc, "5.2 UFO 报表设计器", level=2)
    add_paragraph_styled(doc, "UFO 报表设计器是供专业财务人员进行高阶报表开发的利器，提供类似于 Excel 的在线单元格网格编辑界面：")
    add_bullet_styled(doc, "单元格支持设置常规格式、数值格式、百分比及对齐方式，支持合并单元格。", "格式定义")
    add_bullet_styled(doc, "：在单元格中可输入取数公式。公式格式符合中国财务人员使用习惯：", "财务取数公式")
    add_paragraph_styled(doc, "  - 科目期末余额: LFS(科目编码, 方向, 年, 月) 或 QC() / QM()", bold_prefix="    - ")
    add_paragraph_styled(doc, "  - 科目当期发生额: LFS_FS(科目编码, 借/贷, 年, 月)", bold_prefix="    - ")
    add_paragraph_styled(doc, "  - 单元格四则运算: A1 + B1 - SUM(C1:C10)", bold_prefix="    - ")
    add_bullet_styled(doc, "报表设计完成后，必须锁定公式并保存为正式模板版本，随后可按月一键点击“计算”，系统将自动从数据库读取总账数据，瞬间填满报表并输出计算结果。", "计算与生成")
    
    add_heading_styled(doc, "5.3 预算管理", level=2)
    add_paragraph_styled(doc, "支持科目维度和辅助核算维度的年度与月度预算编制。编制后可以上传至系统，在后续财务报表中输出“预算执行情况表”，比对各科目的“预算数”与“实际发生数”，计算预算达成率和费用超支率，辅助企业日常财务管控。")
    
    add_heading_styled(doc, "5.4 经营多维分析与看板", level=2)
    add_bullet_styled(doc, "：提供图形化看板，直观展示收入/费用趋势图、资产负债结构饼图、现金流瀑布图，并自动计算企业关键财务指标，如流动比率、速动比率、资产负债率、净资产收益率（ROE）等。", "1. 财务 KPI 看板")
    add_bullet_styled(doc, "：分析引擎支持对辅助核算数据进行多维透视。例如，用户可以交叉查询“某项目在某部门下发生的差旅费明细”，或“某供应商在特定期间的采购额趋势”，支持按部门、项目、合作伙伴等多维度进行业绩和成本对比。", "2. 多维透视分析 (Dimensional Analysis)")
    add_bullet_styled(doc, "：支持为常用报表查询条件创建书签，并可配置定时任务。例如，设定“每月5日自动生成上月利润表并导出为 PDF 发送至指定邮箱”。", "3. 书签与定时报表")
    
    # --- 8. 运维与常见问题 ---
    add_heading_styled(doc, "第六章：日常运维与常见问题解答", level=1)
    
    add_heading_styled(doc, "6.1 常用核对性检查清单 (Reconciliation Checklist)", level=2)
    add_paragraph_styled(doc, "为确保系统数据一致性，建议财务人员定期进行以下三项核对工作：")
    add_bullet_styled(doc, "在凭证过账前，运行“试算平衡表”确认本期借贷方发生额合计数是否相等。", "1. 科目试算平衡核对")
    add_bullet_styled(doc, "在期末结账前，进入“库存管理” -> “库存对账”，检查库存明细账的存货价值（如原料、产成品）与总账中对应资产科目（1201/1403/1405等）的余额是否一致。若有差异，需追查是否有手工录入了未挂存货往来的成本凭证，或者领料结转成本凭证未正常过账。", "2. 库存明细账与总账核对")
    add_bullet_styled(doc, "进入“固定资产” -> “固资对账”，核对固定资产卡片原值合计数、累计折旧合计数是否与总账“固定资产”科目（1601）和“累计折旧”科目（1602）的余额完全相等。若不相等，需检查资产原值变动是否已正确生成凭证，或凭证是否已被删除/篡改（系统有防篡改机制，应检查是否有未过账凭证）。", "3. 固定资产卡片与总账核对")
    
    add_heading_styled(doc, "6.2 常见错误及解决方案", level=2)
    
    t_err = create_table_styled(doc, ["错误提示/现象", "可能原因", "解决步骤"], [2.5, 2.0, 3.0])
    add_table_row_styled(t_err, [
        "提示：ACCOUNT_SET_ACCESS_DENIED (403)",
        "当前登录用户未被授予目标账套的访问权限，或者在多账套切换时会话失效。",
        "联系系统管理员在“用户与角色权限”中为该用户增加目标账套的授权，或者尝试退出重新登录。"
    ], is_even=False)
    add_table_row_styled(t_err, [
        "凭证保存时提示：科目要求辅助核算，但分录行未填报",
        "该会计科目在设置中开启了诸如“要求供应商”或“要求部门”的辅助核算标志，但制单时分录行相应辅助字段为空。",
        "检查分录行对应的辅助下拉框，选择对应的核算项；或者修改科目设置（如果该科目确实不需要辅助核算）。"
    ], is_even=True)
    add_table_row_styled(t_err, [
        "凭证保存时提示：会计期间已锁定/关闭",
        "凭证日期对应的会计月份已经关账，或者该期间已被永久锁定。",
        "若期间仅为关闭，且确需补录，可由主管在“会计期间管理”中重新开启该期间；若已锁定，则必须在当前开启的期间内重新制单。"
    ], is_even=False)
    add_table_row_styled(t_err, [
        "成本分摊时提示：存在零数量残值差异",
        "存货账面数量已出完归零，但由于四舍五入或历史单价微差，账面仍有残余金额。",
        "运行系统“残值调整结转”功能，系统会自动生成一笔微调凭证，将残余金额计入当期损益（如管理费用），使金额归零。"
    ], is_even=True)
    
    doc.add_paragraph() # 空行隔开
    
    # 底部说明
    add_paragraph_styled(doc, "--- 手册正文结束 ---", space_after=12)
    
    # 保存路径
    output_path = "docs/AIS_ERP_Suite_User_Manual_Phase1-5.docx"
    doc.save(output_path)
    print(f"User manual generated successfully at: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    generate_manual()
